"""Phase 3a smoke test — case intake backend.

Run from `backend/` with the venv active:

    python tools/smoke_test_phase3.py

Exercises (without needing a running uvicorn server or Celery worker):
  - POST /cases            -> creates case + case_details
  - PATCH /cases/{case_id} -> persists tenant + intake details
  - GET /cases/{case_id}   -> reflects persisted data
  - tools.deadline_calculator, tools.sos_lookup, tools.foundry_iq (direct calls)
  - POST /cases/{case_id}/documents (lease .docx) + POST /cases/{case_id}/submit
  - agents.intake_agent._run(...) run directly (bypassing Celery), then DB
    assertions: case status, deadline/days_overdue, SOS verification,
    agent_session completion, lease_parse_results row, and findings written
    to `case_kb_documents` (Postgres)
  - services.progress_bus publish/subscribe round trip
  - api.routes.chat._handle_chat_message (direct call with a fake websocket)

Cleanup deletes the test case's `lease_parse_results` row(s) (no ON DELETE
CASCADE from `cases`), then the `cases` row (cascades everything else,
including `case_kb_documents`).

Requires AZURE_*, SUPABASE_*, and CELERY_* settings in `backend/.env`
(same as the running app), plus a reachable Redis instance for
`progress_bus` and the Celery broker.
"""

import asyncio
import io
import sys
import time
import uuid
from datetime import date, timedelta
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document as DocxDocument
from httpx import ASGITransport, AsyncClient
from jose import jwt
from sqlalchemy import delete, select

from agents import intake_agent
from api.routes import chat
from config import settings
from database import async_session_factory
from main import app
from models.agent_session import AgentSession
from models.case import Case
from models.case_detail import CaseDetailsSecurityDeposit
from models.case_kb_document import CaseKBDocument
from models.conversation import ConversationMessage
from models.lease_parse import LeaseParseResult
from services import progress_bus
from tools import deadline_calculator, foundry_iq, sos_lookup

LANDLORD_NAME = "GREYSTAR RS GROUP, LLC"

results: list[tuple[str, bool, str]] = []


def check(name: str, ok: bool, detail: str = "") -> None:
    results.append((name, ok, detail))
    status = "PASS" if ok else "FAIL"
    print(f"[{status}] {name}" + (f" - {detail}" if detail else ""))


def make_token() -> str:
    return jwt.encode(
        {
            "sub": "00000000-0000-0000-0000-000000000001",
            "email": "smoke-test@example.com",
            "aud": "authenticated",
            "exp": int(time.time()) + 3600,
        },
        settings.supabase_jwt_secret,
        algorithm="HS256",
    )


class FakeWebSocket:
    """Minimal stand-in for `fastapi.WebSocket` — records sent JSON frames."""

    def __init__(self) -> None:
        self.sent: list[dict] = []

    async def send_json(self, data: dict) -> None:
        self.sent.append(data)


def build_lease_docx() -> bytes:
    doc = DocxDocument()
    doc.add_paragraph("RESIDENTIAL LEASE AGREEMENT")
    doc.add_paragraph("Tenant: Smoke Test Tenant")
    doc.add_paragraph(f"Landlord: {LANDLORD_NAME}")
    doc.add_paragraph("Property Address: 123 Main St, Houston, TX 77002")
    doc.add_paragraph("Lease Start Date: 2024-01-01")
    doc.add_paragraph("Lease End Date: 2024-12-31")
    doc.add_paragraph("Security Deposit: $1,500.00")
    doc.add_paragraph("The Tenant must provide 30 days written notice before vacating the property.")
    doc.add_paragraph("Pets are not allowed without prior written consent from the Landlord.")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


async def run_checks() -> str | None:
    """Run all checks in a single asyncio.run() / event loop. Returns the
    created case_id (for cleanup), or None if case creation itself failed.
    """
    token = make_token()
    headers = {"Authorization": f"Bearer {token}"}
    case_id: str | None = None
    session_id: str | None = None

    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        # --- Step 1: create case ---
        r = await client.post("/cases", headers=headers)
        body = r.json() if r.status_code == 201 else {}
        check("POST /cases -> 201", r.status_code == 201, f"got {r.status_code} {body}")
        if r.status_code != 201:
            return None

        case_id = body["id"]
        check("POST /cases sets status=intake", body.get("status") == "intake", f"got {body.get('status')}")

        # --- Step 2: PATCH intake details, then GET to verify persistence ---
        move_out = date.today() - timedelta(days=60)
        patch_body = {
            "state": "TX",
            "county": "Harris",
            "tenant": {
                "full_legal_name": "Smoke Test Tenant",
                "address": "123 Main St, Houston, TX 77002",
                "phone": "+15555550123",
            },
            "details": {
                "property_address": "123 Main St, Houston, TX 77002",
                "property_state": "TX",
                "property_county": "Harris",
                "property_type": "apartment",
                "landlord_name_as_entered": LANDLORD_NAME,
                "landlord_address": "465 Meeting St Ste 500, Charleston, SC 29403",
                "deposit_amount": 1500,
                "move_in_date": (move_out - timedelta(days=365)).isoformat(),
                "move_out_date": move_out.isoformat(),
                "keys_returned_date": move_out.isoformat(),
                "forwarding_address": "456 Oak St, Houston, TX 77003",
                "forwarding_address_proof": True,
                "landlord_communication": "none",
            },
        }
        r = await client.patch(f"/cases/{case_id}", headers=headers, json=patch_body)
        check("PATCH /cases/{case_id} -> 200", r.status_code == 200, f"got {r.status_code} {r.text}")

        r = await client.get(f"/cases/{case_id}", headers=headers)
        body = r.json() if r.status_code == 200 else {}
        check("GET /cases/{case_id} -> 200", r.status_code == 200, f"got {r.status_code}")

        details_out = body.get("details", {})
        check(
            "GET reflects landlord_name_as_entered",
            details_out.get("landlord_name_as_entered") == LANDLORD_NAME,
            f"got {details_out.get('landlord_name_as_entered')!r}",
        )
        check(
            "GET reflects deposit_amount",
            details_out.get("deposit_amount") == 1500,
            f"got {details_out.get('deposit_amount')!r}",
        )
        parties = body.get("parties", [])
        tenant_party = next((p for p in parties if p.get("role") == "tenant"), None)
        check(
            "GET reflects tenant party",
            tenant_party is not None and tenant_party.get("full_legal_name") == "Smoke Test Tenant",
            f"got {tenant_party!r}",
        )

        # --- Step 3: tools in isolation ---
        dc = deadline_calculator.calculate_deadline("TX", move_out, move_out)
        check(
            "deadline_calculator deadline_date == move_out + 30d",
            dc["deadline_date"] == move_out + timedelta(days=30),
            f"got {dc}",
        )
        check("deadline_calculator days_overdue == 30", dc["days_overdue"] == 30, f"got {dc['days_overdue']}")
        check("deadline_calculator violation_confirmed is True", dc["violation_confirmed"] is True, f"got {dc['violation_confirmed']}")

        try:
            sos_result = await sos_lookup.lookup_entity("TX", LANDLORD_NAME)
            check("sos_lookup.lookup_entity returns 'verified' key", "verified" in sos_result, f"got {sos_result}")
        except Exception as exc:  # noqa: BLE001
            check("sos_lookup.lookup_entity returns 'verified' key", False, f"error: {exc}")

        try:
            chunks = await foundry_iq.query_knowledge_base(
                settings.foundry_kb_state_law, "security deposit return deadline", category="TX", top=3
            )
            check("foundry_iq state-law KB query returns results", len(chunks) > 0, f"got {len(chunks)} chunks")
        except Exception as exc:  # noqa: BLE001
            check("foundry_iq state-law KB query returns results", False, f"error: {exc}")

        # --- Step 4: lease upload, submit, run intake_agent directly ---
        lease_bytes = build_lease_docx()
        files = {
            "file": (
                "lease.docx",
                lease_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }
        r = await client.post(f"/cases/{case_id}/documents", headers=headers, files=files, data={"type": "lease"})
        body = r.json() if r.status_code == 201 else {}
        check("POST /cases/{case_id}/documents -> 201", r.status_code == 201, f"got {r.status_code} {body}")

        r = await client.post(f"/cases/{case_id}/submit", headers=headers)
        body = r.json() if r.status_code == 200 else {}
        check("POST /cases/{case_id}/submit -> 200", r.status_code == 200, f"got {r.status_code} {body}")
        check("submit returns status=researching", body.get("status") == "researching", f"got {body.get('status')}")
        session_id = body.get("session_id")

    if session_id:
        try:
            await intake_agent._run(case_id, session_id)
            check("intake_agent._run completes without raising", True)
        except Exception as exc:  # noqa: BLE001
            check("intake_agent._run completes without raising", False, f"error: {exc}")

    async with async_session_factory() as db:
        case = await db.get(Case, uuid.UUID(case_id))
        check("case.status == assessment", case is not None and case.status == "assessment", f"got {case.status if case else None}")

        case_details = (
            await db.execute(select(CaseDetailsSecurityDeposit).where(CaseDetailsSecurityDeposit.case_id == uuid.UUID(case_id)))
        ).scalar_one_or_none()
        check(
            "details.deadline_date populated",
            case_details is not None and case_details.deadline_date is not None,
            f"got {case_details.deadline_date if case_details else None}",
        )
        check(
            "details.days_overdue == 30",
            case_details is not None and case_details.days_overdue == 30,
            f"got {case_details.days_overdue if case_details else None}",
        )
        check(
            "details.landlord_sos_verified is True",
            case_details is not None and case_details.landlord_sos_verified is True,
            f"got {case_details.landlord_sos_verified if case_details else None}",
        )

        if session_id:
            session = await db.get(AgentSession, uuid.UUID(session_id))
            check(
                "agent_session.status == completed",
                session is not None and session.status == "completed",
                f"got {session.status if session else None}",
            )

        lease_result = (
            await db.execute(select(LeaseParseResult).where(LeaseParseResult.case_id == uuid.UUID(case_id)))
        ).scalar_one_or_none()
        check("lease_parse_results row exists", lease_result is not None)

    # --- intake_agent findings landed in case_kb_documents (Postgres) ---
    async with async_session_factory() as db:
        case_kb_rows = (
            await db.execute(select(CaseKBDocument).where(CaseKBDocument.case_id == uuid.UUID(case_id)))
        ).scalars().all()
        check("case_kb_documents has findings for this case", len(case_kb_rows) > 0, f"got {len(case_kb_rows)} row(s)")

    # --- Step 5: progress_bus publish/subscribe round trip ---
    try:
        events: list[dict] = []

        async def collect() -> None:
            async for event in progress_bus.subscribe(case_id):
                events.append(event)
                return

        task = asyncio.create_task(collect())
        await asyncio.sleep(0.3)  # let the subscriber connect before we publish
        await progress_bus.publish(case_id, {"tool": "smoke_test", "status": "complete"})
        await asyncio.wait_for(task, timeout=10)
        check(
            "progress_bus publish/subscribe round trip",
            bool(events) and events[0].get("tool") == "smoke_test",
            f"got {events}",
        )
    except Exception as exc:  # noqa: BLE001
        check("progress_bus publish/subscribe round trip", False, f"error: {exc}")

    # --- Step 5b: chat._handle_chat_message direct call ---
    try:
        ws = FakeWebSocket()
        async with async_session_factory() as db:
            await asyncio.wait_for(chat._handle_chat_message(ws, db, case_id, "When is my deposit deadline?"), timeout=300)
        event_types = {event.get("type") for event in ws.sent}
        check(
            "chat._handle_chat_message streams tokens and completes",
            "token" in event_types and "done" in event_types,
            f"got event types {event_types}",
        )

        async with async_session_factory() as db:
            assistant_msg = (
                await db.execute(
                    select(ConversationMessage).where(
                        ConversationMessage.case_id == uuid.UUID(case_id), ConversationMessage.role == "assistant"
                    )
                )
            ).scalars().first()
            check("assistant ConversationMessage persisted", assistant_msg is not None)
    except Exception as exc:  # noqa: BLE001
        check("chat._handle_chat_message streams tokens and completes", False, f"error: {exc}")

    return case_id


async def cleanup(case_id: str | None) -> None:
    if case_id is None:
        return

    async with async_session_factory() as db:
        # lease_parse_results has no ON DELETE CASCADE from cases/documents —
        # delete it explicitly before deleting the case.
        await db.execute(delete(LeaseParseResult).where(LeaseParseResult.case_id == uuid.UUID(case_id)))
        case = await db.get(Case, uuid.UUID(case_id))
        if case is not None:
            await db.delete(case)
        await db.commit()


async def run_async_checks() -> None:
    case_id = await run_checks()
    await cleanup(case_id)


def main() -> int:
    asyncio.run(run_async_checks())

    failed = [name for name, ok, _ in results if not ok]
    print()
    print(f"{len(results) - len(failed)}/{len(results)} checks passed")
    if failed:
        print("FAILED:")
        for name in failed:
            print(f"  - {name}")
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
